#!/usr/bin/env python3
"""
Build a response-tracker docx from the referee reports.
Under each referee comment insert a shaded "Response" box with:
  status (ADDRESSED / OPEN / DEFERRED), what was done, location, commit.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from copy import deepcopy

SRC = "/Users/peymansh/MIT Dropbox/Peyman Shahidi/GitHub/hot_towel_revision/referee_reports/Hot Towel - AEJ Micro Revision.docx"
DST = "/Users/peymansh/MIT Dropbox/Peyman Shahidi/GitHub/hot_towel_revision/referee_reports/Hot Towel - AEJ Micro Revision - Response Tracker.docx"

# ------------ Response content keyed by line index in source document ----------
# Keys are the zero-based paragraph indices from the extracted text. Each value is
# a dict with: status, what, location, commit (all optional; at least status+what).

RESPONSES = {
    # Editor's letter (Key goals 1. Model section) - items 14-18 are the sub-bullets a/b/c/d
    14: {"status": "ADDRESSED",
         "what": "Theory section fully restructured: body now contains only the formal model setup, FOCs, and comparative statics; exposition and intuition are presented in separated passages. All proofs, extensions, robustness, and second-signal machinery are in the appendix.",
         "location": "§III.A–E (body); Appendices A (Proofs), B (Information Revelation), C (Model Extensions)",
         "commit": "9d8fcc6 (theory rewrite merge); c3a8cf4 (Corollary 2 proof relocation)"},
    15: {"status": "ADDRESSED",
         "what": "Defensive/hedging passages about generality and alternative modeling choices removed from the body. Alternative production-function discussion moved to Appendix C.",
         "location": "§III body (cleaned); Appendix C",
         "commit": "9d8fcc6"},
    16: {"status": "PARTIALLY ADDRESSED — open items remain",
         "what": "ψ₂ paradox reconciled with a two-pronged response: (i) wage-mechanism clarification at end of §III.C explaining why wage effects on a worker-firm type pair survive the ψ₂→1 limit (pool locked by (m, s₁) before s₂); (ii) IC-based justification for 'almost precise, not perfectly precise' in §III.E / Appendix B.1. OPEN — precise definition of 'large market', truth-telling placement vs. Prop 1, clarification of u' boundedness, role of s₁ still awaiting PK/coauthor call (see Issue #1 residual checklist).",
         "location": "§III.C tail paragraph; §III.E; Appendix B.1 (\\label{online_app:imprecise_second_signal})",
         "commit": "1183eb0; 3f3b5c0; b74b7a0"},
    17: {"status": "ADDRESSED (structural)",
         "what": "Body + main appendix reorganized to be self-contained for Proposition 1. Online Appendix renamed to 'Appendix' throughout (per AEJ Micro's in-print vs. supplementary distinction). Proofs ordered properly: No-mass-point lemma + proof in Appendix A, then Corollary 1, Corollary 2, Proposition 2 (no-talk), Proposition 1 (comparison). Supplementary Empirical Appendix kept separate.",
         "location": "Appendices A/B/C + Supplementary Empirical Appendix",
         "commit": "207d38f (Online Appendix → Appendix sweep)"},

    # Editor Additional comments (list items 23-35)
    23: {"status": "OPEN — John Horton",
         "what": "Needs institutional knowledge of how wages are actually set on Upwork / whether bargaining occurs off-platform. Routed to coauthor.",
         "location": "N/A (response letter item)",
         "commit": "Issue #5 (open, pinged @johnjosephhorton)"},
    24: {"status": "PARTIALLY ADDRESSED — awaiting PK",
         "what": "Kartik (2009) and Abeler et al. (2019) citations added to the bibliography. Issue #4 §2c open: the framing reconciliation between intro ('plausible behavioral assumption') and theory body ('imposed assumption justified by Prop 3') requires PK/Horton steer before committing.",
         "location": "hot_towel.bib (refs added); framing decision pending",
         "commit": "207d38f (bib additions); Issue #4 §2c open"},
    26: {"status": "PARTIALLY ADDRESSED",
         "what": "p.3 'persists from partial to full market roll-out' rewritten for clarity. Citation style cleanup + mid-sentence Horton (2010) reposition + removing 'recent' from 2007-paper references: OPEN pending PK's positioning paragraphs (Issue #3 framing).",
         "location": "p.3 rewrite applied; p.2 citation cleanup pending",
         "commit": "5da3999 (p.3 rewrite)"},
    27: {"status": "ADDRESSED",
         "what": "Rewrote the 'persists from partial to full market roll-out' sentence as: 'our predictions hold whether cheap talk is introduced to a fraction of firms (as in our experimental variation) or to the entire market (as in the platform-wide roll-out that followed the experiment).'",
         "location": "hot_towel.tex line 204",
         "commit": "5da3999"},
    28: {"status": "OPEN — Issue #1 residual",
         "what": "p.8 minor items (repeated 'salient'; standard point (2); rephrasing first part of point (3)) not yet touched. Held for PK's theory-body pass.",
         "location": "§III.A body (top of theory section)",
         "commit": "Issue #1 open"},
    29: {"status": "OPEN — Issue #1 residual",
         "what": "p.9 items not yet touched: a-nonnegative clarification, v_T subscript first-use definition, δ→p rename, 'in principle firms could have…' wording. Held for PK's theory-body pass.",
         "location": "§III.A setup",
         "commit": "Issue #1 open"},
    30: {"status": "OPEN — Issue #1 residual",
         "what": "p.10 'more difficult jobs and worker disutility' clarification not yet touched.",
         "location": "§III.A setup",
         "commit": "Issue #1 open"},
    31: {"status": "OPEN — Issue #1 residual",
         "what": "p.11 'large market' precise definition; footnote 8 citations trim; 'can reject'→'rejects'. 'Can reject'→'rejects' swept in typo pass; other two still pending.",
         "location": "§III",
         "commit": "9c14d38 (partial — 'can reject' fix); Issue #1 residual"},
    32: {"status": "ADDRESSED (typo items)",
         "what": "'message off-equilibrium'→'message off the equilibrium path'; footnote 11 bracket; ':=' consistency — all swept in manuscript-wide typo pass.",
         "location": "§III",
         "commit": "9c14d38 (closes Issue #7)"},
    33: {"status": "ADDRESSED",
         "what": "'Meaningful mass points' clarified; repetition in first sentence of Proposition 1 trimmed.",
         "location": "§III (Proposition 1 walk-through)",
         "commit": "9c14d38"},
    34: {"status": "OPEN — Issue #1 residual",
         "what": "'Generically fails' formality not yet addressed.",
         "location": "§III.E",
         "commit": "Issue #1 open"},
    35: {"status": "ADDRESSED",
         "what": "Footnote 13 'Medium' → 'I' (Intermediate) consistency applied in typo pass.",
         "location": "§III",
         "commit": "9c14d38"},

    # R1 smaller points
    58: {"status": "OPEN — John Horton",
         "what": "Ambiguous 'survey' reference on p.18. Needs Horton to supply the citation or specify which survey is meant before we can edit the text.",
         "location": "Section IV.A (Empirical Context)",
         "commit": "Issue #6 item 2 (open, waiting for Horton)"},
    60: {"status": "ADDRESSED",
         "what": "Typo: 'compared to comparable' rephrased in typo pass.",
         "location": "N/A (find/replace)",
         "commit": "9c14d38"},
    61: {"status": "ADDRESSED",
         "what": "Typo: 'less job applications' → 'fewer job applications'.",
         "location": "N/A",
         "commit": "9c14d38"},
    62: {"status": "ADDRESSED",
         "what": "Typo: 'not specify what message is' reworded.",
         "location": "N/A",
         "commit": "9c14d38"},
    63: {"status": "ADDRESSED",
         "what": "Typo: 'possibly become a possibly' deduplicated.",
         "location": "N/A",
         "commit": "9c14d38"},
    64: {"status": "ADDRESSED",
         "what": "Typo: 'worker differ' → 'workers differ'.",
         "location": "N/A",
         "commit": "9c14d38"},
    65: {"status": "ADDRESSED",
         "what": "Typo: 'statics continues' → 'statics continue'.",
         "location": "N/A",
         "commit": "9c14d38"},
    66: {"status": "ADDRESSED",
         "what": "Typo: broken syntax 'Denote the expected utility they can generate in other markets be' rewritten.",
         "location": "N/A",
         "commit": "9c14d38"},
    67: {"status": "ADDRESSED",
         "what": "Typo: 'due in part to employers consider' → 'due in part to employers considering' (p.28).",
         "location": "hot_towel.tex line 1023",
         "commit": "9c14d38"},

    # R2 §1 Clarity - items 82-91
    82: {"status": "ADDRESSED",
         "what": "Model section rewritten to stand alone in the body: formal model + FOCs + comparative statics in the body; extensions, proofs, and second-signal machinery in the appendix.",
         "location": "§III.A–E body; Appendix A–C",
         "commit": "9d8fcc6"},
    87: {"status": "ADDRESSED",
         "what": "Defensive writing about generality and alternative payoff functions removed from the body; those discussions now live in Appendix C.",
         "location": "§III.A body; Appendix C",
         "commit": "9d8fcc6"},
    88: {"status": "OPEN — Issue #1 residual",
         "what": "'Multiple exogenous signals' introduction (p.10 bottom) and outside-options paragraph still appear without proper setup. Awaiting PK pass.",
         "location": "§III.A",
         "commit": "Issue #1 open"},
    89: {"status": "OPEN — Issue #1 residual",
         "what": "Key assumptions 'mass of workers', 'large market', Poisson-at-each-firm not yet formalized. Awaiting PK pass.",
         "location": "§III.A setup",
         "commit": "Issue #1 open"},
    90: {"status": "OPEN — PK call",
         "what": "Whether to move the Truth-telling subsection (§III.E) earlier in the model section is a structural decision deferred to PK. Currently sits after Proposition 1.",
         "location": "§III.E",
         "commit": "Issue #1 residual (structural)"},
    91: {"status": "ADDRESSED (structure)",
         "what": "Proofs now consolidated in Appendix A in proper order (lemma → Corollary 1 → Corollary 2 → Proposition 2 → Proposition 1). 'Online Appendix' renamed to 'Appendix' to align with AEJ Micro's in-print/supplementary terminology. Supplementary Empirical Appendix is separate.",
         "location": "Appendix A",
         "commit": "c3a8cf4; 207d38f"},

    # R2 §2 Truthtelling
    95: {"status": "ADDRESSED",
         "what": "Added 3-sentence paragraph at end of §III.E emphasizing that downstream predictions on sorting, wages, and efficiency are derived under truth-telling — their empirical confirmation constitutes a second, independent anchor for the assumption alongside Proposition 3's IC argument.",
         "location": "§III.E end, after the rookie/expert bridge paragraph",
         "commit": "9235bc1"},
    96: {"status": "OPEN — awaiting framing call (Issue #4 §2b)",
         "what": "Audit of pooling across Figure 1 (pooled), Figure 4 (not pooled), §V.E (pooled) not yet performed. The rule to apply consistently depends on the framing choice in Issue #4 §2c.",
         "location": "Figures 1, 4 and §V.E",
         "commit": "Issue #4 §2b open"},
    97: {"status": "OPEN — awaiting PK/Horton call (Issue #4 §2c)",
         "what": "Reconciling intro ('plausible behavioral assumption') with theory ('imposed assumption + IC justification'). Routed to PK or Horton before finalizing wording; affects intro narrative and the Kartik/Abeler citation placement.",
         "location": "Intro + §III.E",
         "commit": "Issue #4 §2c open"},

    # R2 §3 Second signal paradox
    99: {"status": "ADDRESSED",
         "what": "Added wage-mechanism paragraph at end of §III.C explaining that wages depend on ā_{s₁} which is pinned down by (m, s₁) before s₂ is observed; the pool is locked at the application stage, so the second signal resolves firm type at the bidding stage without altering the competition.",
         "location": "§III.C tail paragraph",
         "commit": "1183eb0; 3f3b5c0"},
    100: {"status": "ADDRESSED",
          "what": "Explicit IC-based reason for 'almost precise, not perfectly precise' discussed in new subsection 'Why We View the Second Signal as Imprecise' (Appendix B.1): if ψ₂ = 1 exactly, an L-firm deviating to m=H attracts high-ability workers who upon learning the true type bid low, and the deviation becomes profitable; IC requires ψ₂ < 1 but close to 1.",
          "location": "Appendix B.1 (\\label{online_app:imprecise_second_signal})",
          "commit": "b74b7a0 (label fix adding the subsection break)"},

    # R2 smaller comments (102-110)
    102: {"status": "OPEN — Issue #3",
          "what": "Citation style cleanup (p.2 and p.6 Tadelis & Zettelmeyer). Held pending PK's positioning paragraphs.",
          "location": "Intro, §II Literature",
          "commit": "Issue #3 open"},
    103: {"status": "OPEN — Horton/Kircher",
          "what": "First sentence on p.25 ('This simplification of the model…') not yet rewritten; proposed revision drafted and awaiting coauthor sign-off.",
          "location": "§V (around line 898)",
          "commit": "Issue #6 item 3 open"},
    104: {"status": "ADDRESSED",
          "what": "Changed 'prior experience' to 'prior cumulative earnings' on p.26.",
          "location": "hot_towel.tex line 928",
          "commit": "5da3999"},
    105: {"status": "VERIFIED — no change needed",
          "what": "Panel order in Figure 5 ('signal_effects_experience.pdf') verified via pdftotext to already match the text (earnings top, hours-worked bottom). R2's flag appears stale.",
          "location": "Figure 5 + surrounding text",
          "commit": "No change (verified via pdftotext sweep)"},
    106: {"status": "ADDRESSED",
          "what": "Proposition 1 terminology standardized: uniformly uses 'part' throughout. 'Medium'→'I' swept in same pass.",
          "location": "Proposition 1 statement + cross-references",
          "commit": "9c14d38"},
    107: {"status": "OPEN — Horton",
          "what": "'-Y% decrease' phrasing derives from negative values in auto-generated LaTeX parameter macros emitted by the R pipeline. Cleanest fix: wrap the relevant values in abs() inside analysis/params_*.R. Routed to Horton.",
          "location": "analysis/params_*.R",
          "commit": "Issue #9 open (pinged @johnjosephhorton)"},
    108: {"status": "ADDRESSED",
          "what": "Extra ')' after Figure 3 reference removed (line 1016); 'due in part to employers consider' fixed in typo pass.",
          "location": "hot_towel.tex line 1016 (stray ')'); line 1023 (consider→considering)",
          "commit": "5da3999 (stray ')'); 9c14d38 (verb fix)"},
    109: {"status": "OPEN — Horton (if pursuable)",
          "what": "p.34 surplus-vs-wage framing. Requires Horton to confirm whether outside-options data are usable. Minor item per R2.",
          "location": "Section VI",
          "commit": "Response letter item (defer)"},
    110: {"status": "OPEN — Issue #1 residual",
          "what": "p.10 u' bounded vs. unbounded clarification not yet touched. Awaiting PK pass.",
          "location": "§III.A setup",
          "commit": "Issue #1 open"},

    # R3 framing items (117-123)
    117: {"status": "OPEN — coauthors (Issue #3)",
          "what": "Adding Backus, Blake & Tadelis (2019, JPE) reference alongside Tadelis & Zettelmeyer (2015). Needs PK-level positioning paragraph.",
          "location": "§II Literature",
          "commit": "Issue #3 open"},
    118: {"status": "OPEN — coauthors (Issue #3)",
          "what": "Distinguishing Tadelis & Zettelmeyer (2015)'s linkage principle from our IC cheap talk. Needs PK-level framing.",
          "location": "§II Literature",
          "commit": "Issue #3 open"},
    119: {"status": "OPEN — coauthors (Issue #3)",
          "what": "Explicit paragraph connecting our empirical evidence to Backus, Blake & Tadelis (2019)'s framework. R3's main framing ask. Needs PK/Horton.",
          "location": "§II Literature + intro",
          "commit": "Issue #3 open"},
    120: {"status": "OPEN — Peyman (Issue #3 mechanical)",
          "what": "Spell out 'market' vs. 'marketplace' distinction (Roth 2018) more clearly on p.18. Held pending PK's framing direction so language aligns.",
          "location": "§IV.A (Empirical Context)",
          "commit": "Issue #3 Peyman-drafts"},
    121: {"status": "OPEN — coauthors (Issue #3)",
          "what": "Frame probabilistic revelation arm as the garbling device in Backus et al. (2019) and Ambrus, Chaney & Salitskiy (2018, QE). Needs PK-level paragraph.",
          "location": "§IV.A + §II Literature",
          "commit": "Issue #3 open"},
    122: {"status": "ADDRESSED",
          "what": "R3 said 'I did not find these results surprising' for §V.A (Informative Messages). Subsection moved to Empirical Appendix; a two-paragraph preemptive-defense summary was left in the body that frames the pooling question and reports the χ² null, pointing to the appendix for detail.",
          "location": "§V opening + Empirical Appendix \\label{online_app:informative_messages}",
          "commit": "c0ad729"},
    123: {"status": "DECLINED (per R3's own suggestion)",
          "what": "Structural estimation / market-design exercise. R3 explicitly said 'I would actively oppose requiring it for a revision'; kept §V.C in place as a signpost but did not pursue the structural exercise.",
          "location": "No change (§V.C intact)",
          "commit": "Decision noted in response letter"},
    142: {"status": "ADDRESSED",
          "what": "Editor's marginal note 'Maybe move V.A. to Appendix?' — §V.A (Informative Messages) has been moved to the Empirical Appendix.",
          "location": "Empirical Appendix",
          "commit": "c0ad729"},
}

# ------------ Build the document ----------

doc = Document()

# Set default font
from docx.oxml.ns import qn
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_response_box(doc, entry):
    """Insert a single-cell table acting as a shaded box with the response content."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False

    # Indent the whole box a bit
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), '540')  # ~0.375 inch indent
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)

    # Set table width
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), '8500')
    tblW.set(qn('w:type'), 'dxa')

    cell = table.cell(0, 0)
    # Shaded background (light yellow/cream)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF4D6')  # light amber/cream
    tcPr.append(shd)
    # Borders
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'D4A017')
        tcBorders.append(b)
    tcPr.append(tcBorders)

    # Clear any default paragraph in the cell
    cell.text = ''
    p0 = cell.paragraphs[0]

    # Status line — bold, colored by status
    status = entry.get('status', '—')
    if status.startswith('ADDRESSED') or status.startswith('VERIFIED'):
        color = RGBColor(0x1F, 0x7A, 0x1F)  # green
    elif status.startswith('PARTIALLY'):
        color = RGBColor(0xB5, 0x7E, 0x00)  # amber
    elif status.startswith('DECLINED'):
        color = RGBColor(0x6A, 0x4C, 0x93)  # purple
    else:  # OPEN
        color = RGBColor(0xB0, 0x00, 0x1A)  # red

    run = p0.add_run(f'Response: {status}')
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(10)

    # What
    what_para = cell.add_paragraph()
    r = what_para.add_run('What: ')
    r.bold = True
    r.font.size = Pt(10)
    r2 = what_para.add_run(entry.get('what', ''))
    r2.font.size = Pt(10)

    # Location
    if entry.get('location'):
        loc_para = cell.add_paragraph()
        r = loc_para.add_run('Location: ')
        r.bold = True
        r.font.size = Pt(10)
        r2 = loc_para.add_run(entry['location'])
        r2.font.size = Pt(10)

    # Commit
    if entry.get('commit'):
        c_para = cell.add_paragraph()
        r = c_para.add_run('Commit / Issue: ')
        r.bold = True
        r.font.size = Pt(10)
        r2 = c_para.add_run(entry['commit'])
        r2.font.size = Pt(10)

# Header with purpose
title = doc.add_heading('Hot Towel — AEJ Micro Revision Response Tracker', level=0)
intro = doc.add_paragraph()
intro.add_run(
    'This document reproduces the editor\u2019s letter and the three referee reports '
    'verbatim. Under each comment, an inline response box documents the status of '
    'our revision: what was done, where in the paper the fix lives, and the commit '
    'that implemented it. OPEN items note who needs to act and why the item is '
    'deferred. Color key: '
).font.size = Pt(10)
kr = intro.add_run('green ')
kr.bold = True; kr.font.color.rgb = RGBColor(0x1F, 0x7A, 0x1F); kr.font.size = Pt(10)
intro.add_run('= addressed / verified; ').font.size = Pt(10)
ar = intro.add_run('amber ')
ar.bold = True; ar.font.color.rgb = RGBColor(0xB5, 0x7E, 0x00); ar.font.size = Pt(10)
intro.add_run('= partially addressed, open items remain; ').font.size = Pt(10)
rr = intro.add_run('red ')
rr.bold = True; rr.font.color.rgb = RGBColor(0xB0, 0x00, 0x1A); rr.font.size = Pt(10)
intro.add_run('= open, awaiting coauthor input or further work; ').font.size = Pt(10)
pr = intro.add_run('purple ')
pr.bold = True; pr.font.color.rgb = RGBColor(0x6A, 0x4C, 0x93); pr.font.size = Pt(10)
intro.add_run('= declined per referee\u2019s own suggestion.').font.size = Pt(10)

doc.add_paragraph()

# Load source paragraphs with their style info
src = Document(SRC)
# Build list of (style_name, paragraph) to walk
src_paragraphs = list(src.paragraphs)

# Render each paragraph to new doc, inserting response boxes where keys match
for i, p in enumerate(src_paragraphs):
    txt = p.text
    style_name = p.style.name if p.style else 'Normal'

    # Apply style
    if not txt.strip():
        doc.add_paragraph()
    elif style_name.startswith('Title'):
        doc.add_heading(txt, level=0)
    elif style_name.startswith('Heading 1'):
        doc.add_heading(txt, level=1)
    elif style_name.startswith('Heading 2'):
        doc.add_heading(txt, level=2)
    elif style_name.startswith('Heading 3'):
        doc.add_heading(txt, level=2)  # collapse to level 2 for compactness
    elif style_name.startswith('Heading 4'):
        doc.add_heading(txt, level=3)
    elif style_name.startswith('Subtitle'):
        hp = doc.add_paragraph()
        r = hp.add_run(txt); r.italic = True; r.font.size = Pt(12)
    elif style_name.startswith('List'):
        bullet_para = doc.add_paragraph(style='List Bullet')
        bullet_para.add_run(txt)
    else:
        doc.add_paragraph(txt)

    # If this paragraph has a response, attach the box right after
    if i in RESPONSES:
        add_response_box(doc, RESPONSES[i])

doc.save(DST)
print(f'Wrote {DST}')
